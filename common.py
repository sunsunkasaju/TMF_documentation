from docx import Document
import pandas as pd
from markdown import markdown
from io import StringIO
import csv
import sys
import json
from docx.shared import Pt
import re
# Make the doc global:
doc = Document()


def AddJsonSection(Heading, requestJsonFile, responseJsonFile):
    global doc

    doc.add_heading(Heading, 2)

    if requestJsonFile:
        doc.add_heading("Request",3)

        # Read json file and add contents to the document in a monospaced font
        with open(requestJsonFile, 'r') as f:
            data = json.load(f)
        json_data = json.dumps(data, indent=4)
        paragraph = doc.add_paragraph(json_data)
        paragraph.style = 'NoSpacing'
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = 'Courier New'

    if responseJsonFile:
        doc.add_heading("Response",3)

        if Heading == "POST":
            AddMandatoryFields()
        elif Heading == "PATCH":
            AddPatchableFields()
        # Read json file and add contents to the document in a monospaced font
        with open(responseJsonFile, 'r') as f:
            data = json.load(f)
        json_data = json.dumps(data, indent=4)
        paragraph = doc.add_paragraph(json_data)
        paragraph.style = 'NoSpacing'
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = 'Courier New'

    return

def AddMandatoryFields():
    pass

def AddPatchableFields():
    pass


def parseMarkdown(content):
    # Parse the markdown text and apply formatting
    for line in content.split('\n'):
        paragraph = doc.add_paragraph()
        in_bold = False
        in_italics = False

        # Split the line into words
        for word in line.split(' '):
            # Check for bold or italics markdown
            if '**' in word:
                in_bold = not in_bold                
                word = word.replace('**', '')
            if '__' in word:
                in_italics = not in_italics
                word = word.replace('__', '')

            # Add the word with the appropriate formatting
            run = paragraph.add_run(word + ' ')
            run.bold = in_bold
            run.italic = in_italics
            # Reset bold and italic flags after each word
            in_bold = False
            in_italics = False
            

    # Add subheadings
    for subheading in content.split('\n'):
        if subheading.startswith('## '):
            doc.add_heading(subheading.replace('## ', ''), level=2)
        elif subheading.startswith('### '):
            doc.add_heading(subheading.replace('### ', ''), level=3)

def AddSection(Heading, Top, csvLocations, Bottom):
    global doc

    doc.add_heading(Heading, 0)
    parseMarkdown(Top)

    doc.add_heading("Fields:",3)
    
    # Change csvLocation to array of string and call AddTable function iteratively
    for csv in csvLocations:
    
        # Remove the .csv after filename and then change camel case to separate words
        filename = csv.split('/')[-1].replace('.csv', '')
        filename = re.sub('(.)([A-Z][a-z]+)', r'\1 \2', filename)
        filename = re.sub('([a-z0-9])([A-Z])', r'\1 \2', filename)
    
        # Capitalize all the words in filename
        filename = ' '.join(word.capitalize() for word in filename.split())
        
        doc.add_heading(filename, level=2)
        

        AddTable(csv)

        # Add a space before the table
        doc.add_paragraph()




    # Convert markdown Bottom to word format and add to the document
    html = markdown(Bottom)
    paragraph = doc.add_paragraph(html)

    return

def AddTable(csvLocation):
    
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid' 

    try:
        with open(csvLocation, 'r') as csv_file:
            csv_reader = csv.reader(csv_file)

            for row in csv_reader:
                row_cells = table.add_row().cells
                row_cells[0].text = row[0]  # Field Name
                row_cells[1].text = row[1]  # Description

            # Capitalize and bold the column headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = hdr_cells[0].text.capitalize()
            hdr_cells[1].text = hdr_cells[1].text.capitalize()
            hdr_cells[0].paragraphs[0].runs[0].bold = True
            hdr_cells[1].paragraphs[0].runs[0].bold = True

               
    except FileNotFoundError:
        print(f"{csvLocation} CSV files not found.")
        sys.exit(1)



